"""
generate_docx_fixtures.py
Generates .docx fixture files for all 24 IKAM benchmark cases.
Each case gets 1–2 Word documents that are internally consistent with idea.md
and embed intentional contradictions where the chaos level calls for them.

Run:  python3 generate_docx_fixtures.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = os.path.dirname(os.path.abspath(__file__))


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p

def body(doc, text):
    doc.add_paragraph(text)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.text = text
    return p

def table_2col(doc, rows):
    """rows = list of (label, value) tuples"""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = value
    doc.add_paragraph("")

def save(doc, case_id, filename):
    path = os.path.join(BASE, case_id, filename)
    doc.save(path)
    print(f"  → {case_id}/{filename}")


# ---------------------------------------------------------------------------
# s-consulting-v01  |  Birchline Insights
# ---------------------------------------------------------------------------

def gen_s_consulting_v01():
    case = "s-consulting-v01"

    # Doc 1: Engagement scoping memo (surfaces scope contradiction)
    doc = Document()
    h1(doc, "Birchline Insights — Engagement Scoping Memo")
    body(doc, "Date: January 6, 2026")
    body(doc, "Project: Redwood Freight Co. — Network Optimization")
    body(doc, "Prepared by: Avery Hart, Principal")
    doc.add_paragraph("")

    h2(doc, "Engagement Summary")
    body(doc, (
        "This memo confirms the agreed scope of work between Birchline Insights and "
        "Redwood Freight Co. following the contract kickoff call on January 5, 2026."
    ))

    h2(doc, "Agreed Scope")
    body(doc, (
        "Per the executed Statement of Work (SOW) dated December 18, 2025, "
        "this engagement is a 6-week discovery and diagnostic only. "
        "Deliverables include a current-state network map, a root-cause analysis "
        "of routing inefficiencies, and a set of prioritized recommendations. "
        "No implementation work is included in scope."
    ))

    h2(doc, "Key Dates")
    table_2col(doc, [
        ("Kickoff", "January 6, 2026"),
        ("Discovery complete", "February 14, 2026"),
        ("Draft deliverable deck", "February 21, 2026"),
        ("Final deliverable", "February 28, 2026"),
    ])

    h2(doc, "Budget")
    body(doc, (
        "Estimated effort: 130–150 hours across the Birchline team. "
        "Fixed fee per SOW: $28,500. "
        "No change orders are anticipated at this time."
    ))

    h2(doc, "Team")
    table_2col(doc, [
        ("Principal", "Avery Hart"),
        ("Consultant", "Jordan Ng"),
        ("Analyst", "Casey Moore"),
    ])

    h2(doc, "Notes")
    body(doc, (
        "Client requested preliminary feasibility discussion on route optimization "
        "software implementation during kickoff call. This is noted but is not in scope. "
        "Any implementation work would require a separate SOW and additional budget."
    ))
    save(doc, case, "engagement-scoping-memo-2026-01-06.docx")

    # Doc 2: Internal hours reconciliation note (surfaces hours contradiction)
    doc = Document()
    h1(doc, "Birchline Insights — Hours & Invoice Reconciliation Note")
    body(doc, "Date: February 3, 2026")
    body(doc, "Prepared by: Priya Desai, Fractional Finance")
    doc.add_paragraph("")

    h2(doc, "Purpose")
    body(doc, (
        "This note documents the reconciliation between tracked hours in the timesheet "
        "system and hours billed on Invoice #INV-2026-01 for the Redwood Freight engagement."
    ))

    h2(doc, "Hours Summary")
    table_2col(doc, [
        ("Timesheet total (Jan)", "148 hours"),
        ("Hours invoiced", "132 hours"),
        ("Difference", "16 hours"),
        ("Write-off reason", "Non-billable ramp / internal review"),
    ])

    h2(doc, "Detail")
    body(doc, (
        "Avery Hart logged 8 hours of business development calls with Redwood that were "
        "excluded from the invoice per firm policy (BD time is non-billable). "
        "Casey Moore's 6 hours of internal QA review of deliverable slides were also "
        "excluded. Jordan Ng's 2 hours of admin/scheduling are written off."
    ))
    body(doc, (
        "The invoice reflects 132 billable hours. No dispute with client is anticipated. "
        "Final invoice total: $28,500 fixed fee (hours-based rate card not applied)."
    ))

    h2(doc, "Open Items")
    bullet(doc, "Confirm Redwood PO number for invoice.")
    bullet(doc, "Send final deliverable deck before invoice payment terms begin (Net 30).")
    save(doc, case, "hours-reconciliation-2026-02-03.docx")


# ---------------------------------------------------------------------------
# s-construction-v01  |  Alder Ridge Builders
# ---------------------------------------------------------------------------

def gen_s_construction_v01():
    case = "s-construction-v01"

    doc = Document()
    h1(doc, "Alder Ridge Builders — Subcontractor Agreement Checklist")
    body(doc, "Prepared by: Brooke Kim, Project Coordinator")
    body(doc, "Date: February 2026")
    doc.add_paragraph("")

    h2(doc, "Purpose")
    body(doc, (
        "This checklist documents the pre-mobilization requirements for subcontractors "
        "on all active Alder Ridge projects. Use this document to verify compliance "
        "before issuing a Notice to Proceed."
    ))

    h2(doc, "Required Documents")
    table_2col(doc, [
        ("Item", "Status"),
        ("Signed subcontract agreement", "Required before NTP"),
        ("Certificate of insurance (COI)", "Required — Alder Ridge as additional insured"),
        ("W-9 on file", "Required for payment"),
        ("License verification (OR CCB)", "Required"),
        ("Safety plan acknowledgment", "Required"),
        ("Lien waiver (conditional, on first pay)", "Required at first billing"),
    ])

    h2(doc, "Active Projects — Subcontractor Summary (Feb 2026)")
    body(doc, "Cedar Street ADU (Client: Miller family)")
    table_2col(doc, [
        ("Electrical", "Sunrise Electric — COI on file, CCB verified"),
        ("Plumbing", "Pacific Plumbing — COI on file, CCB verified"),
        ("HVAC", "NW Comfort Systems — W-9 pending"),
    ])
    body(doc, "Morningside Office TI (Client: Harmon & Associates)")
    table_2col(doc, [
        ("Drywall", "Cascade Drywall — all docs complete"),
        ("Flooring", "Atlas Floors — COI renewal needed (exp. Feb 28)"),
        ("Electrical", "Sunrise Electric — COI on file"),
    ])

    h2(doc, "Notes")
    body(doc, (
        "Alex Rivera: Atlas Floors COI expired Feb 28 — do not release payment until "
        "renewed COI is on file. Follow up by Feb 20."
    ))
    save(doc, case, "subcontractor-checklist-2026-02.docx")


# ---------------------------------------------------------------------------
# s-healthcare-clinic-v01  |  Pinecrest Family Clinic
# ---------------------------------------------------------------------------

def gen_s_healthcare_clinic_v01():
    case = "s-healthcare-clinic-v01"

    doc = Document()
    h1(doc, "Pinecrest Family Clinic — Monthly Operations Summary")
    body(doc, "Month: January 2026")
    body(doc, "Prepared by: Jordan Kim, Office Manager")
    doc.add_paragraph("")

    h2(doc, "Volume")
    table_2col(doc, [
        ("Total visits", "412"),
        ("New patients", "34"),
        ("Cancellations", "28"),
        ("No-shows", "19"),
        ("No-show rate", "4.6%"),
    ])

    h2(doc, "Billing Snapshot")
    body(doc, (
        "Collections for January are expected to close at approximately $82,000 — "
        "slightly below the $85,000 forecast due to a processing delay from "
        "one of our larger payers (Blue Shield batch held for January cycle). "
        "ClearPath Billing has confirmed the batch will release in February."
    ))
    table_2col(doc, [
        ("Gross charges", "$118,400"),
        ("Expected collections (Jan)", "$82,000"),
        ("Variance vs forecast", "-$3,000 (insurer delay)"),
        ("Outstanding AR >90 days", "$6,200"),
    ])

    h2(doc, "Staffing")
    body(doc, "No staffing changes in January. Dr. Vargas was out Jan 14–15 (conference); "
         "Maya Singh handled patient scheduling coverage with no complaints.")

    h2(doc, "Open Items")
    bullet(doc, "HIPAA refresher training due for Maya Singh by Feb 28.")
    bullet(doc, "Confirm Blue Shield payment posting with ClearPath by Feb 10.")
    bullet(doc, "Schedule quarterly patient satisfaction review — target March.")
    save(doc, case, "ops-summary-2026-01.docx")


# ---------------------------------------------------------------------------
# s-local-retail-v01  |  Bramble & Bitters
# ---------------------------------------------------------------------------

def gen_s_local_retail_v01():
    case = "s-local-retail-v01"

    # Doc 1: Holiday recap (surfaces revenue contradiction)
    doc = Document()
    h1(doc, "Bramble & Bitters — Holiday Season Recap")
    body(doc, "Period: November – December 2025")
    body(doc, "Prepared by: Maya Chen, Owner/GM")
    doc.add_paragraph("")

    h2(doc, "Revenue")
    body(doc, (
        "The 2025 holiday season was our strongest since opening. "
        "Total revenue for Nov–Dec came in at $68,400, a 22% increase over 2024."
    ))
    table_2col(doc, [
        ("In-store sales", "$41,200"),
        ("Online / shipped orders", "$14,800"),
        ("Corporate gift bundles", "$12,400"),
        ("Total", "$68,400"),
    ])

    h2(doc, "Note on Corporate Gifts")
    body(doc, (
        "This was the first year we formally tracked corporate gift bundles as a separate "
        "line. North Ledger LLC books these under 'retail — other' so the bookkeeping "
        "total may show slightly different category splits. The top-line total of $68,400 "
        "is agreed and reconciled."
    ))

    h2(doc, "Top Sellers")
    bullet(doc, "Figgy & Co Apricot Jam 3-pack")
    bullet(doc, "Black Ember Hot Sauce gift set")
    bullet(doc, "Casa Luna Chocolate assortment")
    bullet(doc, "Tea + Honey bundle")

    h2(doc, "What Worked")
    bullet(doc, "Instagram gift-guide posts drove 3× normal DM volume.")
    bullet(doc, "Corporate bundle pricing was well-received — repeat from 6 clients.")
    bullet(doc, "Seasonal subscription box sold out in first 48 hours.")

    h2(doc, "What to Improve")
    bullet(doc, "Inventory management: ran out of top hot sauce SKU Dec 18.")
    bullet(doc, "Online checkout: 4 abandoned cart reports; need simpler flow.")
    bullet(doc, "Holiday staffing: Jordan was stretched — hire seasonal earlier in 2026.")
    save(doc, case, "holiday-recap-2025.docx")

    # Doc 2: Spring pitch (surfaces naming/version contradiction)
    doc = Document()
    h1(doc, "Bramble & Bitters — Spring 2026 Promo Brief")
    body(doc, "Date: January 28, 2026")
    body(doc, "Prepared by: Talia Singh, Marketing Contractor")
    doc.add_paragraph("")

    h2(doc, "Campaign Name")
    body(doc, (
        "Working title: \"Spring Awakening Drop\" (Instagram/TikTok handle: #BrambleSpring26). "
        "Note: internal planning doc from December called this the \"Seasonal Reset\" campaign. "
        "We are aligning on Spring Awakening Drop as the external-facing name going forward."
    ))

    h2(doc, "Timing")
    table_2col(doc, [
        ("Promo window", "March 15 – April 30, 2026"),
        ("Launch post", "March 15 (Instagram)"),
        ("Email announce", "March 14 (day before)"),
        ("New SKUs landing", "March 10 (in store), March 12 (online)"),
    ])

    h2(doc, "Revenue Target")
    body(doc, (
        "Maya's goal: $14,000 in incremental March revenue vs. a typical $10,500 March. "
        "Stretch: $16,000 if the subscription box restock sells out again."
    ))

    h2(doc, "Key Tactics")
    bullet(doc, "Instagram: 3 posts/week, Reels on unboxing new SKUs.")
    bullet(doc, "Email: 2 sends (announce + reminder).")
    bullet(doc, "In-store: window display refresh, end-cap for new spring items.")
    bullet(doc, "Potential TikTok: 1 recipe-style video (hot sauce).")

    h2(doc, "Open Items")
    bullet(doc, "Confirm new SKU list with Maya by Feb 5.")
    bullet(doc, "Draft email copy by Feb 20.")
    bullet(doc, "Finalize hashtag list: #BrambleSpring26 or #BrambleBitters?")
    save(doc, case, "spring-promo-brief-2026.docx")


# ---------------------------------------------------------------------------
# s-manufacturing-v01  |  Ridgeway Metal Works
# ---------------------------------------------------------------------------

def gen_s_manufacturing_v01():
    case = "s-manufacturing-v01"

    doc = Document()
    h1(doc, "Ridgeway Metal Works — Q1 2026 Ops Brief")
    body(doc, "Prepared by: Zoe Park, Office/Operations")
    body(doc, "Date: February 2026")
    doc.add_paragraph("")

    h2(doc, "Shop Capacity & Backlog")
    table_2col(doc, [
        ("Active jobs", "8"),
        ("Quoted / pending", "3"),
        ("Available capacity", "~60 hrs/week remaining"),
        ("Lead time (new orders)", "3–4 weeks"),
    ])

    h2(doc, "Material Costs")
    body(doc, (
        "Steel prices rose approximately 7% in January. Several active quotes were issued "
        "at the prior price point. Hank Miller has approved a surcharge option for "
        "quotes older than 30 days; affected customers will be contacted."
    ))
    body(doc, (
        "Note: The sample bracket run quote issued in December used $0.68/lb for "
        "hot-rolled steel. Current purchase price from Metals USA is $0.73/lb. "
        "The margin on that job is thinner than planned."
    ))

    h2(doc, "Quality")
    body(doc, "No rejections in January. One rework item on the Benson Tools bracket run "
         "(weld porosity) — caught in-house, corrected, no customer impact.")

    h2(doc, "Upcoming")
    bullet(doc, "Schedule annual saw blade maintenance (February).")
    bullet(doc, "Reorder stainless rod stock — Eli flagged level is low.")
    bullet(doc, "Follow up on outstanding payment from Kowalski Industries (Net 30, overdue).")
    save(doc, case, "ops-brief-2026-q1.docx")


# ---------------------------------------------------------------------------
# s-software-v01  |  PocketRelay
# ---------------------------------------------------------------------------

def gen_s_software_v01():
    case = "s-software-v01"

    # Doc 1: Customer onboarding guide (surfaces pricing contradiction)
    doc = Document()
    h1(doc, "PocketRelay — Customer Onboarding Guide")
    body(doc, "Version: January 2026")
    body(doc, "Prepared by: Maya Torres, Customer Success")
    doc.add_paragraph("")

    h2(doc, "Welcome to PocketRelay")
    body(doc, (
        "PocketRelay is a lightweight intake and routing platform built for small and "
        "mid-sized operations teams. This guide walks you through account setup, "
        "your first routing rule, and common support questions."
    ))

    h2(doc, "Current Plans")
    table_2col(doc, [
        ("Starter", "$49/month — up to 3 users, 500 submissions/mo"),
        ("Growth", "$99/month — up to 10 users, unlimited submissions"),
        ("Enterprise", "Custom pricing — contact sales"),
    ])
    body(doc, (
        "Note: A small number of customers who signed up before June 2025 are on "
        "legacy pricing. If you believe you are on a legacy plan, please contact "
        "support at support@pocketrelay.io to confirm your rate."
    ))

    h2(doc, "Getting Started")
    body(doc, "Step 1: Log in at app.pocketrelay.io with your registered email.")
    body(doc, "Step 2: Navigate to Settings → Routing Rules → Add Rule.")
    body(doc, "Step 3: Define your intake form fields and destination queue.")
    body(doc, "Step 4: Test your rule with a sample submission.")

    h2(doc, "Support")
    body(doc, "Email: support@pocketrelay.io | Response SLA: 1 business day for Growth; "
         "2 business days for Starter.")
    save(doc, case, "customer-onboarding-guide-2026.docx")

    # Doc 2: Churn + metrics note (surfaces churn definition contradiction)
    doc = Document()
    h1(doc, "PocketRelay — Monthly Metrics Note")
    body(doc, "Month: January 2026")
    body(doc, "Prepared by: Dev Patel, Fractional Finance")
    doc.add_paragraph("")

    h2(doc, "MRR Snapshot")
    table_2col(doc, [
        ("MRR (start of Jan)", "$18,400"),
        ("New MRR", "$1,200"),
        ("Churned MRR (cancellations + downgrades)", "-$620"),
        ("MRR (end of Jan)", "$19,000 (est.)"),
    ])

    h2(doc, "Churn Definition — Finance View")
    body(doc, (
        "For finance reporting, 'churn' includes both full cancellations and plan downgrades "
        "(e.g., Growth → Starter). Net revenue churn for January: -$620 / $18,400 = 3.4%. "
        "This is within normal range; no corrective action required."
    ))

    h2(doc, "Note on Support Reporting")
    body(doc, (
        "Maya's support dashboard uses 'cancellations' to mean accounts that fully closed "
        "their subscription. She reported 2 cancellations in January. Finance sees a total "
        "of 4 churned events when including 2 downgrades. Both are correct — they measure "
        "different things. Riley is aware and will align definitions before the next board update."
    ))
    save(doc, case, "monthly-metrics-note-2026-01.docx")


# ---------------------------------------------------------------------------
# m-consulting-v01  |  Northlake Advisory Group
# ---------------------------------------------------------------------------

def gen_m_consulting_v01():
    case = "m-consulting-v01"

    doc = Document()
    h1(doc, "Northlake Advisory Group — Engagement Status Report")
    body(doc, "Date: February 10, 2026")
    body(doc, "Prepared by: Casey Nguyen, Engagement Manager")
    doc.add_paragraph("")

    h2(doc, "Active Engagements")
    table_2col(doc, [
        ("Client", "Status"),
        ("Cascade Logistics (Codename: Ironwood)", "In delivery — Week 6 of 12"),
        ("Pacific Health Partners (Codename: Elm)", "Kickoff complete — discovery underway"),
        ("Regent Capital (Codename: Birch)", "Proposal submitted — decision expected Feb 20"),
        ("Thornton Manufacturing", "Retainer active — monthly advisory"),
    ])

    h2(doc, "Cascade Logistics — Utilization Note")
    body(doc, (
        "The practice dashboard shows 78% consultant utilization for January. "
        "However, Morgan Lee spent approximately 30 hours on the Cascade engagement "
        "doing internal model builds that were coded as 'internal development' "
        "rather than client-billable. If reclassified, utilization for January would be "
        "closer to 71%. Elena Park and Finance have been notified; coding will be "
        "corrected going forward."
    ))

    h2(doc, "Key Risks")
    bullet(doc, "Regent Capital decision is delayed — Birch pipeline revenue is at risk for Q1.")
    bullet(doc, "Pacific Health data access issues (IT procurement) may push discovery by 1 week.")
    bullet(doc, "Priya Shah at capacity; no additional engagements before March.")

    h2(doc, "Open Items")
    bullet(doc, "Elena Park to send Regent Capital follow-up by Feb 12.")
    bullet(doc, "Casey Nguyen to update scope memo for Pacific Health by Feb 14.")
    bullet(doc, "Finance to reclass Morgan Lee hours in January time log.")
    save(doc, case, "engagement-status-2026-02-10.docx")


# ---------------------------------------------------------------------------
# m-construction-v01  |  Beaconline Construction Co.
# ---------------------------------------------------------------------------

def gen_m_construction_v01():
    case = "m-construction-v01"

    doc = Document()
    h1(doc, "Beaconline Construction Co. — Project Status Memo")
    body(doc, "Project: Harborview Office Park TI — Suite 400")
    body(doc, "Date: February 17, 2026")
    body(doc, "Prepared by: Sara Ito, Senior PM")
    doc.add_paragraph("")

    h2(doc, "Schedule")
    body(doc, (
        "As of this memo, the project is tracking toward a Substantial Completion "
        "target of March 28, 2026, per the agreed milestone schedule communicated "
        "to the owner in the February 10 weekly meeting."
    ))
    body(doc, (
        "Note: The schedule XLSX (Rev 4, issued February 3) reflects a re-baselined "
        "Substantial Completion of April 5, 2026 due to the delayed delivery of "
        "the custom millwork package. There is a discrepancy between the date "
        "communicated verbally and the current schedule file. Sara Ito to present "
        "updated schedule to owner at the February 24 OAC meeting."
    ))

    h2(doc, "Change Orders")
    table_2col(doc, [
        ("CO #1 — Electrical panel upgrade", "$28,400 — Executed"),
        ("CO #2 — Additional data drops (8)", "$14,200 — Executed"),
        ("CO #3 — Flooring upgrade (owner)", "$31,800 — Executed"),
        ("CO #4 — Millwork revision", "$44,000 — Executed"),
        ("CO #5 — Fire sprinkler relocation", "$Pending — in review"),
        ("CO Log Total (executed)", "$118,400"),
    ])
    body(doc, (
        "Note: The owner's representative referenced a total change order value of "
        "$124,900 in their February 10 email. This figure appears to include CO #5 "
        "(pending, not executed). Beaconline's executed CO total is $118,400."
    ))

    h2(doc, "Budget / Margin")
    body(doc, (
        "PM cost report projects a 6% fee margin at completion. "
        "Accounting WIP worksheet (Controller Priyanka Nair) shows tighter margin "
        "in the 3–4% range based on actual cost-to-date vs. revised budget. "
        "Will reconcile at next monthly owner billing cycle (March 1)."
    ))
    save(doc, case, "project-status-memo-2026-02-17.docx")


# ---------------------------------------------------------------------------
# m-healthcare-clinic-v01  |  Mesa Ridge Urgent & Primary Care
# ---------------------------------------------------------------------------

def gen_m_healthcare_clinic_v01():
    case = "m-healthcare-clinic-v01"

    doc = Document()
    h1(doc, "Mesa Ridge Urgent & Primary Care — January 2026 Operations Report")
    body(doc, "Prepared by: Ryan Patel, Ops Manager")
    body(doc, "Date: February 5, 2026")
    doc.add_paragraph("")

    h2(doc, "Volume Across Sites")
    table_2col(doc, [
        ("Site", "Visits"),
        ("Mesa North", "1,104"),
        ("Mesa Central", "892"),
        ("Ridgeline", "743"),
        ("Total", "2,739"),
    ])

    h2(doc, "No-Show Rate")
    body(doc, (
        "Ops tracking (which excludes late-cancel appointments) shows a network-wide "
        "no-show rate of 5.8% for January. "
        "Site managers at Mesa North and Ridgeline have flagged that their local tracking "
        "includes late cancels (within 2 hours of appointment), which brings their "
        "combined no-show figure to approximately 8.1%. "
        "Elise Grant (billing) uses the higher definition for revenue cycle projections. "
        "Ryan Patel to propose unified definition at February ops meeting."
    ))

    h2(doc, "Staffing")
    body(doc, (
        "All three sites fully staffed. Two MA positions at Mesa Central filled in January "
        "(start dates Jan 6 and Jan 13). Tasha Nguyen to update the staffing model "
        "spreadsheet to reflect new hires."
    ))

    h2(doc, "Billing Snapshot")
    table_2col(doc, [
        ("Network collections (Jan, est.)", "$318,000"),
        ("Finance summary (Jan, preliminary)", "$310,500"),
        ("Variance", "~$7,500 — timing of posted payments"),
    ])
    save(doc, case, "ops-report-2026-01.docx")


# ---------------------------------------------------------------------------
# m-local-retail-v01  |  Driftwood & Oak
# ---------------------------------------------------------------------------

def gen_m_local_retail_v01():
    case = "m-local-retail-v01"

    doc = Document()
    h1(doc, "Driftwood & Oak — Spring 2026 Campaign Brief")
    body(doc, "Prepared by: Dani Park, Marketing")
    body(doc, "Date: February 2026")
    doc.add_paragraph("")

    h2(doc, "Campaign Name")
    body(doc, (
        "This campaign has been referred to as \"Spring Reset\" in the internal promo calendar. "
        "For Instagram and external channels, we are using \"New Season Drop.\" "
        "All customer-facing creative should use New Season Drop."
    ))

    h2(doc, "Revenue Goal")
    body(doc, (
        "Avery's target: $28,000 combined March–April incremental revenue across both locations. "
        "South location (Riley) is projecting $14,500; East location (Taylor) is projecting $13,500."
    ))

    h2(doc, "Key Products")
    bullet(doc, "New spring ceramics line (8 SKUs, handmade, limited run)")
    bullet(doc, "Outdoor entertaining accessories")
    bullet(doc, "Seasonal candle restocks (citrus/herb scents)")

    h2(doc, "Channels")
    table_2col(doc, [
        ("Instagram", "3 posts/week + Stories"),
        ("Email", "2 sends — announce + last-call"),
        ("In-store", "Window refresh, focal table display"),
        ("Wholesale pitch", "2 boutique hotel accounts (Avery managing)"),
    ])

    h2(doc, "Open Items")
    bullet(doc, "Finalize product photography by Feb 22 (Mina Lopez to coordinate).")
    bullet(doc, "Confirm inventory levels for ceramics before launch date.")
    bullet(doc, "Align campaign name with Avery — is 'Spring Reset' still in use anywhere?")
    save(doc, case, "spring-campaign-brief-2026.docx")


# ---------------------------------------------------------------------------
# m-manufacturing-v01  |  Northpoint Packaging & Plastics
# ---------------------------------------------------------------------------

def gen_m_manufacturing_v01():
    case = "m-manufacturing-v01"

    doc = Document()
    h1(doc, "Northpoint Packaging & Plastics — February 2026 Ops Review")
    body(doc, "Prepared by: Simone Grant, GM")
    body(doc, "Date: February 14, 2026")
    doc.add_paragraph("")

    h2(doc, "Production")
    table_2col(doc, [
        ("Units produced (Jan)", "842,000"),
        ("Units planned (Jan)", "900,000"),
        ("Attainment", "93.6%"),
        ("Root cause of gap", "Resin delivery delay (Vendor: MidWest Poly, Jan 9)"),
    ])

    h2(doc, "Scrap / Rework")
    body(doc, (
        "QA reports scrap rate at 1.4% for January (excludes rework; only counts units "
        "that cannot be recovered). "
        "Ops tracking includes rework in the scrap figure, showing 2.9% total defect rate. "
        "Simone Grant to align QA and Ops on a single definition before the Q1 customer "
        "scorecards are issued in April."
    ))

    h2(doc, "Inventory Valuation Note")
    body(doc, (
        "Finance (Jordan Blake) carries resin inventory at standard cost ($0.94/lb for PETG). "
        "Last purchase from MidWest Poly was at $1.02/lb (January). "
        "There is a $0.08/lb variance on approximately 180,000 lbs of current stock, "
        "or roughly $14,400 total. This will be flagged in the Q1 close."
    ))

    h2(doc, "On-Time Delivery")
    body(doc, (
        "Customer scorecard (based on ship date vs. PO due date) shows 94.2% OTD for Q4 2025. "
        "This calculation counts partial shipments as on-time if the primary run shipped "
        "on time. Customers Vertex Packaging and AllFresh Foods have each separately noted "
        "that they count partial shipments as late. "
        "Derek Rao to review definition with top 5 accounts."
    ))
    save(doc, case, "ops-review-2026-02.docx")


# ---------------------------------------------------------------------------
# m-software-v01  |  LatticeOps
# ---------------------------------------------------------------------------

def gen_m_software_v01():
    case = "m-software-v01"

    doc = Document()
    h1(doc, "LatticeOps — Feature Garnet Status Update")
    body(doc, "Date: February 14, 2026")
    body(doc, "Prepared by: Alex Chen, Head of Product")
    doc.add_paragraph("")

    h2(doc, "Feature Overview")
    body(doc, (
        "Feature Garnet is the next major product milestone for LatticeOps — a "
        "workflow automation layer allowing customers to trigger integrations directly "
        "from ops boards. It is the anchor feature for H1 2026 roadmap delivery."
    ))

    h2(doc, "Current Status")
    body(doc, "Status: In Development — Sprint 6 of 9")

    h2(doc, "Ship Date")
    body(doc, (
        "The H1 2026 roadmap deck (presented to board in January) lists Garnet's "
        "target release as April 2026. "
        "Based on the current sprint log (Sprint 6 closed February 7), engineering "
        "has re-forecasted the release to late May 2026 due to the API authentication "
        "rework required after the February 9 incident. "
        "This delay has not yet been communicated to customers."
    ))

    h2(doc, "February 9 Incident — Impact on Garnet")
    body(doc, (
        "The incident triggered a mandatory security review of all new API surface area. "
        "Garnet's integration endpoints are in scope. Sam Rivera estimates 3 additional "
        "sprints of work. The postmortem attributed root cause to vendor API instability, "
        "but internal sprint metrics show a significant backlog spike on the LatticeOps side "
        "that contributed to the degraded error handling. The auth rework addresses both."
    ))

    h2(doc, "KPI Alignment Note")
    body(doc, (
        "Garnet adoption metrics will be measured differently by CS and Product post-launch: "
        "CS will track 'active accounts' using Garnet at least once per week. "
        "Product will track 'weekly active users' across all accounts. "
        "Jamie Ross (CEO) has asked both teams to align on a single primary metric "
        "before the launch announcement."
    ))
    save(doc, case, "feature-garnet-status-2026-02-14.docx")


# ---------------------------------------------------------------------------
# l-consulting-v01  |  Greybridge Partners
# ---------------------------------------------------------------------------

def gen_l_consulting_v01():
    case = "l-consulting-v01"

    doc = Document()
    h1(doc, "Greybridge Partners — January 2026 Practice Performance Summary")
    body(doc, "Prepared by: Finance & Operations")
    body(doc, "Date: February 6, 2026")
    doc.add_paragraph("")

    h2(doc, "Utilization")
    body(doc, (
        "The January practice dashboard (distributed to partners) shows a firm-wide "
        "utilization rate of 78% for client-billable work. "
        "The finance tracker calculates January utilization at 71%, as it excludes "
        "hours coded as 'internal development', 'proposal writing', and 'BD support' "
        "from the denominator. "
        "Both figures are accurate within their respective definitions. "
        "The Partners agreed at the January 30 leadership meeting to adopt the finance "
        "tracker definition (71%) as the official utilization KPI going forward, "
        "effective February reporting."
    ))

    h2(doc, "Revenue")
    table_2col(doc, [
        ("January recognized revenue", "$1,840,000"),
        ("January billed", "$1,760,000"),
        ("Unbilled (WIP)", "$80,000"),
        ("Pipeline (Q1 weighted)", "$3,200,000"),
    ])

    h2(doc, "Active Engagements — Highlights")
    bullet(doc, "Hartwell Logistics: Phase 2 kickoff Feb 3. On track.")
    bullet(doc, "Meridian Health: Diagnostic complete. Readout Feb 12.")
    bullet(doc, "Blackstone Foods: Final report delivered. Closeout pending.")
    bullet(doc, "Fenway Capital: Proposal submitted Jan 28. Decision by Feb 15.")

    h2(doc, "Risk Items")
    bullet(doc, "Two senior consultants at capacity — no capacity for new large engagements until March.")
    bullet(doc, "Meridian Health scope creep risk: client requesting additional deliverables outside SOW.")
    save(doc, case, "practice-performance-2026-01.docx")


# ---------------------------------------------------------------------------
# l-construction-v01  |  Ironwood Ridge Constructors
# ---------------------------------------------------------------------------

def gen_l_construction_v01():
    case = "l-construction-v01"

    doc = Document()
    h1(doc, "Ironwood Ridge Constructors — Project Milestone Report")
    body(doc, "Project: Mesa Commerce Center — Building B")
    body(doc, "Date: February 20, 2026")
    body(doc, "Prepared by: Project Management")
    doc.add_paragraph("")

    h2(doc, "Milestone Summary")
    table_2col(doc, [
        ("Structural steel complete", "January 14, 2026 — Complete"),
        ("MEP rough-in", "February 28, 2026 — On track"),
        ("Exterior cladding", "March 30, 2026 — On track"),
        ("Substantial Completion", "October 18, 2026 — Owner update"),
        ("Final Completion", "November 5, 2026 — Per internal schedule"),
    ])

    h2(doc, "Substantial Completion — Schedule Note")
    body(doc, (
        "The Substantial Completion date of October 18, 2026 was communicated to the "
        "owner in the January 28 project update email. "
        "The current internal schedule (Rev 6, February 10) reflects a Substantial "
        "Completion date of November 5, 2026, incorporating a weather delay buffer "
        "and a pending inspection cycle for the rooftop mechanical units. "
        "The project team will present the revised schedule to the owner at the "
        "February 25 OAC meeting and seek written acceptance."
    ))

    h2(doc, "Change Orders")
    body(doc, (
        "Active change order log total: $2,180,000 across 11 executed COs. "
        "Three additional COs are in review (CO #12–14), totaling approximately $340,000. "
        "Owner representative has verbally approved CO #12; written approval pending."
    ))

    h2(doc, "Open Issues")
    bullet(doc, "Glazing subcontractor lead time — confirm delivery date by Feb 28.")
    bullet(doc, "Fire marshal pre-inspection scheduled for March 10.")
    bullet(doc, "Owner to provide final finish selections for lobby by March 1.")
    save(doc, case, "milestone-report-2026-02-20.docx")


# ---------------------------------------------------------------------------
# l-healthcare-clinic-v01  |  Harborview Care Partners
# ---------------------------------------------------------------------------

def gen_l_healthcare_clinic_v01():
    case = "l-healthcare-clinic-v01"

    doc = Document()
    h1(doc, "Harborview Care Partners — Network Operations Summary")
    body(doc, "Month: January 2026")
    body(doc, "Prepared by: Operations")
    doc.add_paragraph("")

    h2(doc, "Volume Across 6 Sites")
    table_2col(doc, [
        ("Total visits (Jan)", "8,412"),
        ("New patients", "682"),
        ("No-shows (ops definition)", "4.9%"),
        ("No-shows (revenue cycle definition)", "7.1%"),
        ("Late cancels (excluded from ops)", "182 events"),
    ])

    h2(doc, "No-Show Definition Discrepancy")
    body(doc, (
        "Operations dashboard excludes late-cancel appointments (within 4 hours of start) "
        "from the no-show rate calculation. The reported rate for January is 4.9%. "
        "The revenue cycle team includes late-cancels as no-shows for billing projection "
        "purposes, yielding a rate of 7.1%. "
        "Both figures appear in network reporting and have caused confusion with site directors. "
        "Medical Director Dr. Elena Marsh has asked for a unified definition to be adopted "
        "before the Q1 board report in April."
    ))

    h2(doc, "Billing")
    table_2col(doc, [
        ("Gross charges (Jan)", "$2,840,000"),
        ("Expected collections", "$1,920,000"),
        ("Collection rate (target)", "67%"),
        ("Collection rate (actuals est.)", "65.8%"),
    ])

    h2(doc, "Compliance")
    body(doc, "HIPAA refresher training completion: 94% network-wide. "
         "Sites at Marin and Sacramento-East are below 80% — HR to follow up by Feb 28.")
    save(doc, case, "network-ops-summary-2026-01.docx")


# ---------------------------------------------------------------------------
# l-local-retail-v01  |  Juniper & Juno Goods
# ---------------------------------------------------------------------------

def gen_l_local_retail_v01():
    case = "l-local-retail-v01"

    doc = Document()
    h1(doc, "Juniper & Juno Goods — Weekly Ops Report")
    body(doc, "Week of: February 10–16, 2026")
    body(doc, "Prepared by: Store Operations")
    doc.add_paragraph("")

    h2(doc, "Revenue Summary")
    body(doc, (
        "This report uses gross sales figures (before returns and promotional discounts). "
        "Bookkeeping (handled by Oakhurst Ledger) uses net sales for monthly reporting, "
        "which will differ from these weekly totals by approximately 4–6% depending on "
        "return volume."
    ))
    table_2col(doc, [
        ("Silver Lake", "$14,820 gross"),
        ("Echo Park", "$11,340 gross"),
        ("Atwater Village", "$9,760 gross"),
        ("Online", "$4,200 gross"),
        ("Total (gross)", "$40,120"),
    ])

    h2(doc, "Promotions")
    bullet(doc, "15% Valentine's bundle promo ran Feb 10–14 — performed above forecast.")
    bullet(doc, "Spring collection preview in-store only (no online yet).")

    h2(doc, "Inventory")
    body(doc, "Echo Park flagged two SKUs at reorder level: linen napkin set and ceramic mug 4-pack. "
         "Purchasing to replenish by Feb 20.")

    h2(doc, "Staffing")
    body(doc, "Atwater Village short-staffed Feb 12 (call-out) — Silver Lake covered 4-hour gap. "
         "No customer complaints received.")
    save(doc, case, "weekly-ops-report-2026-02-16.docx")


# ---------------------------------------------------------------------------
# l-manufacturing-v01  |  Cascadia Fasteners & Forming
# ---------------------------------------------------------------------------

def gen_l_manufacturing_v01():
    case = "l-manufacturing-v01"

    doc = Document()
    h1(doc, "Cascadia Fasteners & Forming — January 2026 Production Summary")
    body(doc, "Prepared by: Operations")
    body(doc, "Date: February 4, 2026")
    doc.add_paragraph("")

    h2(doc, "Output")
    table_2col(doc, [
        ("Units produced (Jan)", "4,820,000"),
        ("Planned", "5,000,000"),
        ("Attainment", "96.4%"),
        ("Top gap reason", "Scheduled maintenance (Jan 18–19) — planned downtime"),
    ])

    h2(doc, "Inventory Valuation Note")
    body(doc, (
        "Finance carries fastener-grade steel at standard cost ($0.72/lb). "
        "The most recent purchase from Pacific Steel (January 8) was at $0.78/lb. "
        "Current stock: approximately 420,000 lbs — $25,200 valuation variance. "
        "Finance will adjust at Q1 close."
    ))

    h2(doc, "On-Time Delivery")
    table_2col(doc, [
        ("OTD rate (Jan)", "97.1%"),
        ("Calculation basis", "Ship date vs. PO due date; complete orders only"),
        ("3PL transfer accuracy", "99.4%"),
    ])

    h2(doc, "Quality")
    body(doc, "PPM (parts per million defective) for January: 284. Target: <400. "
         "No customer escapes reported.")
    save(doc, case, "production-summary-2026-01.docx")


# ---------------------------------------------------------------------------
# l-software-v01  |  QuillStack Systems
# ---------------------------------------------------------------------------

def gen_l_software_v01():
    case = "l-software-v01"

    doc = Document()
    h1(doc, "QuillStack Systems — Feature Fable Delivery Update")
    body(doc, "Date: February 18, 2026")
    body(doc, "Prepared by: Product Management")
    doc.add_paragraph("")

    h2(doc, "Feature Overview")
    body(doc, (
        "Feature Fable introduces dynamic workflow branching, allowing enterprise customers "
        "to configure conditional routing in QuillStack reports without custom code."
    ))

    h2(doc, "Target Ship Date")
    body(doc, (
        "The Q1 roadmap deck (presented to customers and investors in January 2026) "
        "lists Fable's GA target as May 2026. "
        "The current sprint log (as of Sprint 8 completed Feb 14) forecasts a June 2026 "
        "release due to additional testing requirements for multi-tenant data isolation. "
        "Customer-facing communication will be updated once the engineering team confirms "
        "the revised forecast at the February 24 sprint review."
    ))

    h2(doc, "Status by Workstream")
    table_2col(doc, [
        ("Core branching logic", "Complete"),
        ("UI configuration panel", "In progress — 70%"),
        ("Multi-tenant isolation testing", "Not started — begins Sprint 9"),
        ("Documentation", "In progress"),
        ("Beta customer onboarding", "Scheduled post-GA"),
    ])

    h2(doc, "Risks")
    bullet(doc, "Multi-tenant testing is the critical path item — any slip extends the June target.")
    bullet(doc, "Two beta customers (Harlan Corp, Meridian Trust) expecting May access; will need proactive outreach.")
    save(doc, case, "feature-fable-update-2026-02-18.docx")


# ---------------------------------------------------------------------------
# xl-consulting-v01  |  Alderpoint Strategy & Operations (ASO)
# ---------------------------------------------------------------------------

def gen_xl_consulting_v01():
    case = "xl-consulting-v01"

    doc = Document()
    h1(doc, "Alderpoint Strategy & Operations — Q4 2025 Client Impact Summary")
    body(doc, "Prepared by: Finance & Delivery Operations")
    body(doc, "Date: January 2026")
    doc.add_paragraph("")

    h2(doc, "Firm-Wide Savings Delivered (2025)")
    body(doc, (
        "This document summarizes the annualized savings delivered to ASO clients "
        "in 2025, as reported across two internal sources."
    ))
    table_2col(doc, [
        ("Source", "Figure"),
        ("Exec deck (board presentation)", "$18.4M annualized savings"),
        ("Finance impact model", "$14.9M net savings"),
        ("Variance", "$3.5M"),
    ])

    h2(doc, "Reconciliation Note")
    body(doc, (
        "The executive deck figure ($18.4M) uses the clients' own baseline estimates "
        "as the reference point, and applies gross savings before one-time implementation "
        "costs. The finance model ($14.9M) uses ASO's internally validated baseline "
        "and deducts estimated one-time transformation costs borne by clients. "
        "Both figures are accurate within their respective definitions. "
        "Going forward, all external communications will use the net figure ($14.9M) "
        "per Managing Director guidance."
    ))

    h2(doc, "Top Engagements by Impact")
    table_2col(doc, [
        ("Halloran Industrials (procurement)", "$4.2M net"),
        ("Pacific Transport (network redesign)", "$3.8M net"),
        ("Westfield Health (cost structure)", "$2.9M net"),
        ("Cascade Retail (inventory + ops)", "$2.1M net"),
        ("Other (8 engagements)", "$1.9M net"),
    ])

    h2(doc, "Utilization")
    body(doc, (
        "Firm utilization for Q4 2025: 74% (finance tracker definition). "
        "The practice dashboard shows 80% for the same period, as it includes "
        "proposal and BD time in the denominator differently."
    ))
    save(doc, case, "q4-impact-summary-2025.docx")


# ---------------------------------------------------------------------------
# xl-construction-v01  |  SummitSpan Builders Group
# ---------------------------------------------------------------------------

def gen_xl_construction_v01():
    case = "xl-construction-v01"

    doc = Document()
    h1(doc, "SummitSpan Builders Group — Enterprise Project Summary")
    body(doc, "Report Period: February 2026")
    body(doc, "Prepared by: PMO")
    doc.add_paragraph("")

    h2(doc, "Active Projects — Schedule Highlights")
    table_2col(doc, [
        ("Project", "Substantial Completion"),
        ("Cascade Tower (Portland)", "July 15, 2026 — per PM weekly notes"),
        ("Cascade Tower (Portland)", "August 2, 2026 — per master schedule (re-baselined)"),
        ("Ridgeline Logistics Park", "March 30, 2026 — on track"),
        ("Harbor Commons Multifamily", "September 12, 2026"),
        ("Eastgate Industrial Phase 2", "December 2026 (TBD)"),
    ])

    h2(doc, "Cascade Tower — Schedule Discrepancy")
    body(doc, (
        "The PM weekly notes for Cascade Tower, distributed to the ownership team "
        "in the January 27 project report, reference a Substantial Completion date of "
        "July 15, 2026. "
        "The master schedule (Rev 9, issued February 3) shows August 2, 2026 after "
        "re-baselining for weather delays on the podium level and a revised elevator "
        "procurement timeline. "
        "The PMO will present the revised schedule to the owner at the February 26 "
        "OAC meeting and seek written acknowledgment."
    ))

    h2(doc, "Safety")
    body(doc, "Recordable incident rate (RIR) across active projects: 1.2 (trailing 12 months). "
         "Industry benchmark for commercial GC: ~2.0. No lost-time incidents in January.")

    h2(doc, "Financial Summary (All Projects)")
    table_2col(doc, [
        ("Total backlog (contracted)", "$284M"),
        ("Revenue in progress (Jan)", "$18.4M"),
        ("Gross margin (projected, portfolio)", "8.2%"),
    ])
    save(doc, case, "enterprise-project-summary-2026-02.docx")


# ---------------------------------------------------------------------------
# xl-healthcare-clinic-v01  |  Harborview Family Clinics Network (HFCN)
# ---------------------------------------------------------------------------

def gen_xl_healthcare_clinic_v01():
    case = "xl-healthcare-clinic-v01"

    doc = Document()
    h1(doc, "Harborview Family Clinics Network — January 2026 Network Report")
    body(doc, "Prepared by: Network Operations & Billing")
    body(doc, "Date: February 8, 2026")
    doc.add_paragraph("")

    h2(doc, "Volume — 9 Sites")
    table_2col(doc, [
        ("Total network visits", "21,840"),
        ("New patients", "1,620"),
        ("No-show rate (ops dashboard)", "7.5%"),
        ("No-show rate (site directors)", "10–12%"),
    ])

    h2(doc, "No-Show Definition Note")
    body(doc, (
        "The ops dashboard (maintained centrally) excludes appointments cancelled "
        "with less than 2 hours notice from the no-show calculation — these are "
        "counted as cancellations. The reported rate is 7.5%. "
        "Site directors at 6 of 9 clinics include these late-cancel events as "
        "no-shows in their local tracking, consistent with how their scheduling "
        "software flags them. Their blended estimate is 10–12%. "
        "The Network Medical Director has asked for a standardized definition to be "
        "implemented by April 1, 2026 (ahead of the annual board presentation)."
    ))

    h2(doc, "Billing")
    table_2col(doc, [
        ("Gross charges (Jan)", "$7,820,000"),
        ("Expected collections", "$5,180,000"),
        ("Collection rate (target)", "66%"),
        ("Collection rate (actuals, est.)", "64.2%"),
        ("Variance reason", "Payer delay — Aetna batch held for processing"),
    ])

    h2(doc, "Membership Pilot")
    body(doc, (
        "The membership access pilot (3 clinics, 120 enrolled members) is tracking "
        "at 88% renewal intent on the January survey. Finance is projecting $18,000 "
        "MRR from the pilot by Q2 if enrollment holds."
    ))
    save(doc, case, "network-report-2026-01.docx")


# ---------------------------------------------------------------------------
# xl-local-retail-v01  |  Sunbeam Market Collective
# ---------------------------------------------------------------------------

def gen_xl_local_retail_v01():
    case = "xl-local-retail-v01"

    doc = Document()
    h1(doc, "Sunbeam Market Collective — Weekly Operations Report")
    body(doc, "Week of: February 9–15, 2026")
    body(doc, "Prepared by: Store Operations Team")
    doc.add_paragraph("")

    h2(doc, "Revenue Summary")
    body(doc, (
        "Weekly ops revenue figures use gross sales (before returns, voids, and "
        "promotional discounts). Finance reports net sales monthly. "
        "Expect a 3–6% variance between this report and the monthly finance summary."
    ))
    table_2col(doc, [
        ("Mission (SF)", "$42,100 gross"),
        ("Oakland Grand Ave", "$38,400 gross"),
        ("Berkeley", "$31,200 gross"),
        ("San Jose Santana Row", "$28,800 gross"),
        ("Marin", "$24,100 gross"),
        ("Palo Alto", "$22,600 gross"),
        ("Commissary (wholesale)", "$14,800 gross"),
        ("Total (gross)", "$202,000"),
    ])

    h2(doc, "Waste")
    body(doc, (
        "Network-wide prepared food waste for the week: 8.2% of prepared food production. "
        "Commissary tracking shows 6.1% waste on its production log. "
        "The difference reflects store-level measurement variability — the commissary uses "
        "weight-based tracking; stores use visual/count estimation. "
        "Ops to standardize tracking method by March 1."
    ))

    h2(doc, "Meal Kit Subscription Pilot")
    body(doc, (
        "Week 3 of the meal kit subscription pilot. 214 active subscribers across Mission "
        "and Berkeley pickup. Fulfillment accuracy: 97.2%. "
        "2 customer complaints (missing item; wrong variant) resolved same-day."
    ))
    save(doc, case, "weekly-ops-report-2026-02-15.docx")


# ---------------------------------------------------------------------------
# xl-manufacturing-v01  |  Titan River Components
# ---------------------------------------------------------------------------

def gen_xl_manufacturing_v01():
    case = "xl-manufacturing-v01"

    doc = Document()
    h1(doc, "Titan River Components — February 2026 Executive Operations Report")
    body(doc, "Prepared by: Operations & Finance")
    body(doc, "Date: February 12, 2026")
    doc.add_paragraph("")

    h2(doc, "OEE — Overall Equipment Effectiveness")
    body(doc, (
        "Plant-level dashboards report OEE at 78% across the three facilities for January. "
        "The finance operations review uses a different availability assumption "
        "(scheduled maintenance windows are excluded from available time in plant dashboards "
        "but included in the finance model denominator), yielding an OEE of 72%. "
        "Both figures are presented to the board; the finance figure (72%) is used "
        "for capital allocation planning."
    ))

    h2(doc, "Production — All Plants")
    table_2col(doc, [
        ("Plant A (Detroit)", "2,840,000 units — 97.2% attainment"),
        ("Plant B (Dayton)", "1,920,000 units — 94.1% attainment"),
        ("Plant C (Indianapolis)", "3,110,000 units — 96.8% attainment"),
        ("Total", "7,870,000 units"),
    ])

    h2(doc, "Quality — Scrap Rate")
    body(doc, (
        "Finance and quality report scrap rate at 1.8% (units scrapped vs. produced). "
        "Operations tracks a broader 'total defect rate' that includes reworked units "
        "before reclassification, showing 3.1%. "
        "Customer-facing scorecards use the 1.8% figure."
    ))

    h2(doc, "Inventory")
    body(doc, (
        "Raw material inventory is valued at standard cost in the ERP. "
        "Steel coil (primary input) was purchased at $840/ton in January vs. "
        "the $790/ton standard. Variance of approximately $620,000 across "
        "current stock will be recognized at Q1 close."
    ))
    save(doc, case, "exec-ops-report-2026-02.docx")


# ---------------------------------------------------------------------------
# xl-software-v01  |  Meridian SignalWorks
# ---------------------------------------------------------------------------

def gen_xl_software_v01():
    case = "xl-software-v01"

    doc = Document()
    h1(doc, "Meridian SignalWorks — ARR Reconciliation Memo")
    body(doc, "Date: February 2026")
    body(doc, "Prepared by: Finance")
    doc.add_paragraph("")

    h2(doc, "Purpose")
    body(doc, (
        "This memo reconciles the ARR figures presented in the January board deck "
        "and the current finance model, which differ by approximately $900,000."
    ))

    h2(doc, "ARR Figures")
    table_2col(doc, [
        ("Board deck (January)", "$24.8M ARR"),
        ("Finance model (current)", "$23.9M ARR"),
        ("Difference", "$900,000"),
    ])

    h2(doc, "Explanation")
    body(doc, (
        "The board deck ARR figure ($24.8M) includes $900,000 of revenue from three "
        "annual contracts that were pre-paid in 2025 and are now churned "
        "(customers did not renew for 2026). These were included in the ARR snapshot "
        "because the contracts had not yet expired at the time of the January board meeting. "
        "The finance model ($23.9M) removes these three accounts as they are confirmed "
        "non-renewals, consistent with our ARR policy of excluding known churn "
        "from the forward-looking metric."
    ))

    h2(doc, "Resolution")
    body(doc, (
        "The $23.9M figure should be used in all investor communications, "
        "sales materials, and internal reporting going forward. "
        "The next board update (March) will present ARR at $23.9M with an explanation "
        "of the prior period correction."
    ))

    h2(doc, "Usage-Based Add-On")
    body(doc, (
        "Usage-based add-on revenue of approximately $1.2M (annualized) is NOT "
        "included in either ARR figure above, per company policy of excluding "
        "variable revenue from ARR. This is tracked separately as usage revenue."
    ))
    save(doc, case, "arr-reconciliation-memo-2026-02.docx")

    # Second doc: product status
    doc = Document()
    h1(doc, "Meridian SignalWorks — Product Milestone Update")
    body(doc, "Date: February 2026")
    body(doc, "Prepared by: Product")
    doc.add_paragraph("")

    h2(doc, "Active Milestones")
    table_2col(doc, [
        ("Feature / Release", "Target"),
        ("Atlas v4.0 (EU data residency)", "March 31, 2026"),
        ("Signal Pulse (usage analytics)", "Q2 2026"),
        ("Enterprise SSO revamp", "April 2026 (in progress)"),
        ("Workflow Composer (major)", "H2 2026 TBD"),
    ])

    h2(doc, "EU Data Residency")
    body(doc, (
        "Atlas v4.0 (EU data residency) is the top priority for Q1. "
        "Required for three enterprise customers in the UK/EU expansion pipeline. "
        "Engineering is on track for a March 31 GA. Legal review of updated DPA "
        "language is in progress."
    ))

    h2(doc, "KPI Definitions")
    body(doc, (
        "Following the Q4 KPI audit, Finance and Product have aligned on the following: "
        "ARR excludes usage revenue and known churned annual prepaid. "
        "NRR (net revenue retention) includes expansions and downgrades but excludes "
        "new logos. The next board package will include both ARR and NRR."
    ))
    save(doc, case, "product-milestone-update-2026-02.docx")


# ---------------------------------------------------------------------------
# Entry point — run all generators
# ---------------------------------------------------------------------------

GENERATORS = [
    ("s-consulting-v01", gen_s_consulting_v01),
    ("s-construction-v01", gen_s_construction_v01),
    ("s-healthcare-clinic-v01", gen_s_healthcare_clinic_v01),
    ("s-local-retail-v01", gen_s_local_retail_v01),
    ("s-manufacturing-v01", gen_s_manufacturing_v01),
    ("s-software-v01", gen_s_software_v01),
    ("m-consulting-v01", gen_m_consulting_v01),
    ("m-construction-v01", gen_m_construction_v01),
    ("m-healthcare-clinic-v01", gen_m_healthcare_clinic_v01),
    ("m-local-retail-v01", gen_m_local_retail_v01),
    ("m-manufacturing-v01", gen_m_manufacturing_v01),
    ("m-software-v01", gen_m_software_v01),
    ("l-consulting-v01", gen_l_consulting_v01),
    ("l-construction-v01", gen_l_construction_v01),
    ("l-healthcare-clinic-v01", gen_l_healthcare_clinic_v01),
    ("l-local-retail-v01", gen_l_local_retail_v01),
    ("l-manufacturing-v01", gen_l_manufacturing_v01),
    ("l-software-v01", gen_l_software_v01),
    ("xl-consulting-v01", gen_xl_consulting_v01),
    ("xl-construction-v01", gen_xl_construction_v01),
    ("xl-healthcare-clinic-v01", gen_xl_healthcare_clinic_v01),
    ("xl-local-retail-v01", gen_xl_local_retail_v01),
    ("xl-manufacturing-v01", gen_xl_manufacturing_v01),
    ("xl-software-v01", gen_xl_software_v01),
]

if __name__ == "__main__":
    for case_id, gen_fn in GENERATORS:
        print(f"\n{case_id}")
        gen_fn()
    print("\nDone. All .docx files written.")
